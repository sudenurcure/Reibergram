import os
import datetime
import matplotlib.pyplot as plt
from App import plot_reibergram
from PyQt5.QtWidgets import (
    QApplication,
    QPushButton,
    QWidget,
    QLabel,
    QLineEdit,
    QGridLayout,
    QMessageBox,
)
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QKeyEvent
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

# Constants for document format settings
WORD_FORMAT = "Word (.docx)"


def create_date_folder():
    today = datetime.date.today()
    folder_name = today.strftime("%Y-%m-%d")
    folder_path = os.path.join("Tüm Belgeler", folder_name)

    if not os.path.exists(folder_path):
        os.makedirs(folder_path)

    return folder_path


class DataEntryWindow(QWidget):
    labels = [
        "Adı Soyadı:",
        "Yaşı:",
        "Cinsiyeti:",
        "Örnek Numarası:",
        "QIgG:",
        "QAlb:",
    ]

    def __init__(self):
        super().__init__()

        self.init_ui()

    def init_ui(self):
        layout = QGridLayout()
        self.input_widgets = {}  # Store QLineEdit widgets in a dictionary
        self.add_input(layout, "Adı Soyadı:")
        self.add_input(layout, "Yaşı:")
        self.add_input(layout, "Cinsiyeti:")
        self.add_input(layout, "Örnek Numarası:")
        self.add_input(layout, "QIgG:")
        self.add_input(layout, "QAlb:")

        self.submit_button = QPushButton("Kaydet")
        self.submit_button.clicked.connect(self.save_data)
        layout.addWidget(self.submit_button, 7, 0, 1, 2)
        self.submit_button.setFixedWidth(120)

        self.reset_button = QPushButton("Temizle")
        self.reset_button.clicked.connect(self.reset_fields)
        layout.addWidget(self.reset_button, 8, 0, 1, 2)
        self.reset_button.setFixedWidth(80)

        # Create a hidden button for the "Enter" key action
        self.hidden_button = QPushButton("HiddenButton")
        self.hidden_button.setHidden(True)
        layout.addWidget(self.hidden_button)

        self.setLayout(layout)

        for widget in self.input_widgets.values():
            widget.installEventFilter(self)

        # Connect the "Enter" key press event to the "Kaydet" button click event
        self.hidden_button.clicked.connect(self.submit_button.click)
        self.submit_button.setDefault(True)  # This makes it respond to the Enter key
        self.setWindowTitle("Hasta Bilgileri")

    def eventFilter(self, obj, event):
        if event.type() == QKeyEvent.KeyPress:
            if event.key() in [Qt.Key_Return, Qt.Key_Down]:
                current_widget = QApplication.focusWidget()
                current_label = None
                for label_text, widget in self.input_widgets.items():
                    if widget == current_widget:
                        current_label = label_text
                        break
                current_index = (
                    self.labels.index(current_label) if current_label else -1
                )
                if 0 <= current_index < len(self.labels) - 1:
                    next_label = self.labels[current_index + 1]
                    next_widget = self.input_widgets.get(next_label)
                    if next_widget:
                        next_widget.setFocus()
                elif current_index == len(self.labels) - 1:
                    self.submit_button.setFocus()

            elif event.key() == Qt.Key_Up:
                current_widget = QApplication.focusWidget()
                current_label = None
                for label_text, widget in self.input_widgets.items():
                    if widget == current_widget:
                        current_label = label_text
                        break
                current_index = (
                    self.labels.index(current_label) if current_label else -1
                )
                if current_index > 0:
                    prev_label = self.labels[current_index - 1]
                    prev_widget = self.input_widgets.get(prev_label)
                    if prev_widget:
                        prev_widget.setFocus()

            elif event.key() == Qt.Key_Enter:
                self.submit_button.click()  # Trigger "Kaydet" button click on Enter key press

        return super().eventFilter(obj, event)

    def add_input(self, layout, label_text):
        label = QLabel(label_text)
        entry = QLineEdit()
        layout.addWidget(
            label, self.labels.index(label_text), 0, Qt.AlignmentFlag.AlignLeft
        )
        layout.addWidget(entry, self.labels.index(label_text), 1)
        self.input_widgets[
            label_text
        ] = entry  # Store the QLineEdit widget in the dictionary

    def reset_fields(self):
        for widget in self.input_widgets.values():
            widget.clear()

    def save_data(self):
        name = self.input_widgets["Adı Soyadı:"].text()
        age = self.input_widgets["Yaşı:"].text()
        gender = self.input_widgets["Cinsiyeti:"].text()
        barcode = self.input_widgets["Örnek Numarası:"].text()
        qigg = self.input_widgets["QIgG:"].text()
        qalb = self.input_widgets["QAlb:"].text()

        # Validate inputs
        if not all([name, age, gender, barcode, qigg, qalb]):
            QMessageBox.warning(
                self,
                "Validasyon Hatası!",
                "Lütfen tüm boşlukları doldurunuz.",
                QMessageBox.Ok,
            )
            return
        sex = ["K", "E", "M", "F"]
        if gender.upper() not in sex:
            QMessageBox.warning(
                self,
                "Validasyon Hatası!",
                "Lütfen geçerli (K/E) bir cinsiyet giriniz.",
                QMessageBox.Ok,
            )
            return

        try:
            age = int(age)
            qigg = float(qigg) / 1000
            qalb = float(qalb) / 1000
        except ValueError:
            QMessageBox.warning(
                self,
                "Validasyon Hatası!",
                "Lütfen geçerli bir Yaş, QIgG veya QAlb değeri giriniz.",
                QMessageBox.Ok,
            )
            return

        # Call the function to generate the Word document with information and the Reibergram plot
        folder_path = create_date_folder()
        self.generate_word(
            qigg, qalb, name.upper(), age, gender.upper(), barcode, folder_path
        )

        # Reset the input fields
        self.reset_fields()
        QMessageBox.information(
            self,
            "Veriler Kaydedildi",
            "Kaydedildi.",
            QMessageBox.Ok,
        )

    def generate_word(self, Qigg, Qalbumin, name, age, gender, barcode, folder_path):
        doc_name = f"{barcode}.docx"
        plot_file = f"{barcode}.png"
        IgA = "IgA.png"
        IgM = "IgM.png"
        doc_path = os.path.join(folder_path, doc_name)

        # Create a Word document
        doc = Document()
        plot_reibergram(Qigg, Qalbumin, barcode)

        # Add collected information to the Word document
        info_text = f"\nAdı Soyadı: {name}\nCinsiyeti, yaşı: {gender}/{age}\nÖrnek No: {barcode}\nRapor Tarihi: {datetime.date.today().strftime('%d.%m.%Y')}"

        table = doc.add_table(rows=3, cols=2)
        table.autofit = False

        # Set the column widths
        table.columns[0].width = Cm(6.43)
        table.columns[1].width = Cm(6.43)

        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if i == 0:
                    if j == 0:
                        # First cell in the first row - add text
                        paragraph = cell.add_paragraph(info_text)
                        paragraph.paragraph_format.line_spacing_rule = (
                            WD_LINE_SPACING.SINGLE
                        )
                        run = paragraph.runs[0]
                        run.font.size = Pt(12)
                        run.font.bold = True
                        run.font.name = "Times New Roman"
                    elif j == 1:
                        # Second cell in the first row - add image
                        text = ["BOS/Serum quotient diagramları \n(Reibergram)"]
                        paragraph = cell.add_paragraph()
                        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph.paragraph_format.line_spacing_rule = (
                            WD_LINE_SPACING.SINGLE
                        )
                        p_run = paragraph.add_run(text)
                        p_run.font.size = Pt(12)
                        p_run.font.bold = True
                        p_run.font.name = "Times New Roman"

                        cell.paragraphs[1].format_alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = cell.paragraphs[1].add_run()
                        run.add_picture(plot_file, width=Cm(6.6), height=Cm(6.4))
                elif i == 1:
                    if j == 0:
                        pass
                    elif j == 1:
                        # Second cell in the second row - add image
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = cell.paragraphs[0].add_run()
                        run.add_picture(IgA, width=Cm(6.6), height=Cm(6.4))
                else:
                    if j == 0:
                        pass
                    else:
                        # Second cell in the third row - add image
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = cell.paragraphs[0].add_run()
                        run.add_picture(IgM, width=Cm(6.6), height=Cm(6.4))

        # Save the Word document
        doc.save(doc_path)

        # Close the plot to free up memory
        plt.close()

        # Delete the plot file
        os.remove(plot_file)
