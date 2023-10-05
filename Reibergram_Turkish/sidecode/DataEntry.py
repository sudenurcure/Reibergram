import os
import datetime
import matplotlib.pyplot as plt
from App import plot_reibergram
from PyQt5.QtWidgets import (
    QApplication,
    QPushButton,
    QWidget,
    QLabel,
    QLineEdit,
    QGridLayout,
    QMessageBox,
    QCheckBox,
)
from PyQt5.QtCore import Qt, QSettings
from PyQt5.QtGui import QKeyEvent
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Pt
from docx.shared import Inches


# Constants for document format settings
WORD_FORMAT = "Word (.docx)"
PDF_FORMAT = "PDF (.pdf)"


def create_date_folder():
    today = datetime.date.today()
    folder_name = today.strftime("%Y-%m-%d")
    folder_path = os.path.join("Tüm Belgeler", folder_name)

    if not os.path.exists(folder_path):
        os.makedirs(folder_path)

    return folder_path


class DataEntryWindow(QWidget):
    labels = [
        "Adı Soyadı:",
        "Yaşı:",
        "Cinsiyeti:",
        "Örnek Numarası:",
        "QIgG:",
        "QAlb:",
    ]

    def __init__(self):
        super().__init__()

        self.init_ui()

    def init_ui(self):
        layout = QGridLayout()
        self.input_widgets = {}  # Store QLineEdit widgets in a dictionary
        self.add_input(layout, "Adı Soyadı:")
        self.add_input(layout, "Yaşı:")
        self.add_input(layout, "Cinsiyeti:")
        self.add_input(layout, "Örnek Numarası:")
        self.add_input(layout, "QIgG:")
        self.add_input(layout, "QAlb:")

        # Add checkboxes for selecting document formats
        self.pdf_checkbox = QCheckBox(PDF_FORMAT)
        self.word_checkbox = QCheckBox(WORD_FORMAT)
        layout.addWidget(self.pdf_checkbox, 6, 1)
        layout.addWidget(self.word_checkbox, 6, 2)

        self.submit_button = QPushButton("Kaydet")
        self.submit_button.clicked.connect(self.save_data)
        layout.addWidget(self.submit_button)

        self.reset_button = QPushButton("Temizle")
        self.reset_button.clicked.connect(self.reset_fields)
        layout.addWidget(self.reset_button)

        # Create a hidden button for the "Enter" key action
        self.hidden_button = QPushButton("HiddenButton")
        self.hidden_button.setHidden(True)
        layout.addWidget(self.hidden_button)

        self.setLayout(layout)

        for widget in self.input_widgets.values():
            widget.installEventFilter(self)

        # Connect the "Enter" key press event to the "Kaydet" button click event
        self.hidden_button.clicked.connect(self.submit_button.click)

        # Load the user's format preference from settings
        settings = QSettings("YourCompany", "YourApp")
        saved_format = settings.value("document_format", "")
        if PDF_FORMAT in saved_format:
            self.pdf_checkbox.setChecked(True)
        if WORD_FORMAT in saved_format:
            self.word_checkbox.setChecked(True)

    def eventFilter(self, obj, event):
        if event.type() == QKeyEvent.KeyPress:
            if event.key() in [Qt.Key_Return, Qt.Key_Down]:
                current_widget = QApplication.focusWidget()
                current_label = None
                for label_text, widget in self.input_widgets.items():
                    if widget == current_widget:
                        current_label = label_text
                        break
                current_index = (
                    self.labels.index(current_label) if current_label else -1
                )
                if 0 <= current_index < len(self.labels) - 1:
                    next_label = self.labels[current_index + 1]
                    next_widget = self.input_widgets.get(next_label)
                    if next_widget:
                        next_widget.setFocus()
                elif current_index == len(self.labels) - 1:
                    self.submit_button.setFocus()

            elif event.key() == Qt.Key_Up:
                current_widget = QApplication.focusWidget()
                current_label = None
                for label_text, widget in self.input_widgets.items():
                    if widget == current_widget:
                        current_label = label_text
                        break
                current_index = (
                    self.labels.index(current_label) if current_label else -1
                )
                if current_index > 0:
                    prev_label = self.labels[current_index - 1]
                    prev_widget = self.input_widgets.get(prev_label)
                    if prev_widget:
                        prev_widget.setFocus()
        return super().eventFilter(obj, event)

    def add_input(self, layout, label_text):
        label = QLabel(label_text)
        entry = QLineEdit()
        layout.addWidget(
            label, self.labels.index(label_text), 0, Qt.AlignmentFlag.AlignLeft
        )
        layout.addWidget(entry, self.labels.index(label_text), 1)
        self.input_widgets[
            label_text
        ] = entry  # Store the QLineEdit widget in the dictionary

    def reset_fields(self):
        for widget in self.input_widgets.values():
            widget.clear()

    def save_data(self):
        # Get the selected formats
        selected_formats = []
        if self.pdf_checkbox.isChecked():
            selected_formats.append(PDF_FORMAT)
        if self.word_checkbox.isChecked():
            selected_formats.append(WORD_FORMAT)

        # Ensure at least one format is selected
        if not selected_formats:
            QMessageBox.warning(
                self,
                "Format Seçimi Hatası!",
                "Lütfen en az bir dosya formatı seçiniz.",
                QMessageBox.Ok,
            )
            return

        # Store the user's format preferences in settings
        settings = QSettings("YourCompany", "YourApp")
        settings.setValue("document_format", " ".join(selected_formats))

        # Get the entered values
        name = self.input_widgets["Adı Soyadı:"].text()
        age = self.input_widgets["Yaşı:"].text()
        gender = self.input_widgets["Cinsiyeti:"].text()
        barcode = self.input_widgets["Örnek Numarası:"].text()
        qigg = self.input_widgets["QIgG:"].text()
        qalb = self.input_widgets["QAlb:"].text()

        # Validate inputs
        if not all([name, age, gender, barcode, qigg, qalb]):
            QMessageBox.warning(
                self,
                "Validasyon Hatası!",
                "Lütfen tüm boşlukları doldurunuz.",
                QMessageBox.Ok,
            )
            return
        sex = ["K", "E", "M", "F"]
        if gender.upper() not in sex:
            QMessageBox.warning(
                self,
                "Validasyon Hatası!",
                "Lütfen geçerli (K/E) bir cinsiyet giriniz.",
                QMessageBox.Ok,
            )
            return

        try:
            age = int(age)
            qigg = float(qigg) / 1000
            qalb = float(qalb) / 1000
        except ValueError:
            QMessageBox.warning(
                self,
                "Validasyon Hatası!",
                "Lütfen geçerli bir Yaş, QIgG veya QAlb değeri giriniz.",
                QMessageBox.Ok,
            )
            return

        # Call the function to generate the selected documents with information and the Reibergram plot
        folder_path = create_date_folder()
        if PDF_FORMAT in selected_formats:
            self.generate_pdf(
                qigg, qalb, name, age, gender.upper(), barcode, folder_path
            )
        if WORD_FORMAT in selected_formats:
            self.generate_word(
                qigg, qalb, name, age, gender.upper(), barcode, folder_path
            )

        # Reset the input fields
        self.reset_fields()
        QMessageBox.information(
            self,
            "Veriler Kaydedildi",
            "Kaydedildi.",
            QMessageBox.Ok,
        )

    def generate_word(self, Qigg, Qalbumin, name, age, gender, barcode, folder_path):
        doc_name = f"{barcode}.docx"
        plot_file = f"{barcode}.png"
        doc_path = os.path.join(folder_path, doc_name)

        # Create a Word document
        doc = Document()

        plot_reibergram(Qigg, Qalbumin, barcode)

        # Add collected information to the Word document
        info_text = f"Adı Soyadı: {name}\nYaşı: {age}\nCinsiyeti: {gender}\nÖrnek Numarası: {barcode}\nRapor Tarihi: {datetime.date.today()}"

        para = doc.add_paragraph()
        run = para.add_run(info_text)
        font = run.font
        font.size = Pt(14)
        font.bold = True

        # Add the plot image to the Word document
        doc.add_picture(
            plot_file, width=Inches(5), height=Inches(5)
        )  # Adjust width and height as needed

        # Save the Word document
        doc.save(doc_path)

        # Close the plot to free up memory
        plt.close()

        # Delete the plot file
        os.remove(plot_file)

    def generate_pdf(self, Qigg, Qalbumin, name, age, gender, barcode, folder_path):
        pdf_name = f"{barcode}.pdf"
        plot_file = f"{barcode}.png"
        pdf_path = os.path.join(folder_path, pdf_name)

        # Create a PDF document
        c = canvas.Canvas(pdf_path, pagesize=letter)

        # Add collected information to the PDF document with line breaks
        info_lines = [
            f"Adı Soyadı: {name}",
            f"Yaşı: {age}",
            f"Cinsiyeti: {gender}",
            f"Örnek Numarası: {barcode}",
            f"Rapor Tarihi: {datetime.date.today()}",
        ]

        x = 100  # X-coordinate for the text
        y = 750  # Initial Y-coordinate for the text

        # Set the font and font size
        font_name = "Helvetica"
        font_size = 12

        for line in info_lines:
            c.setFont(font_name, font_size)
            c.drawString(x, y, line)
            y -= font_size * 1.2  # Adjust the line spacing as needed

        plot_reibergram(Qigg, Qalbumin, barcode)

        # Add the plot image to the PDF document
        c.drawImage(
            plot_file, 100, 50, width=350, height=350
        )  # Adjust position, width, and height as needed

        # Save the PDF document
        c.save()

        # Close the plot to free up memory
        plt.close()

        # Delete the plot file
        os.remove(plot_file)


if __name__ == "__main__":
    import sys

    app = QApplication(sys.argv)
    main_window = DataEntryWindow()
    main_window.show()
    sys.exit(app.exec_())
