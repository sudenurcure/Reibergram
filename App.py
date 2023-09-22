import sys
import os
from PyQt5.QtWidgets import (
    QApplication,
    QMainWindow,
    QPushButton,
    QWidget,
    QVBoxLayout,
    QLabel,
    QLineEdit,
    QGridLayout,
    QMessageBox,
)
from PyQt5.QtCore import Qt, QTimer
from PyQt5.QtGui import QGuiApplication
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from Integrated import *
from docx import Document
from docx.shared import Inches
import datetime
import shutil


def create_date_folder():
    today = datetime.date.today()
    folder_name = today.strftime("%Y-%m-%d")
    folder_path = os.path.join("All Documents", folder_name)

    if not os.path.exists(folder_path):
        os.makedirs(folder_path)

    return folder_path


class DataEntryWindow(QWidget):
    def __init__(self):
        super().__init__()

        self.init_ui()

    def init_ui(self):
        layout = QGridLayout()

        self.input_widgets = {}  # Store QLineEdit widgets in a dictionary

        self.add_input(layout, "Name/Surname:", 0)
        self.add_input(layout, "Age:", 1)
        self.add_input(layout, "Gender:", 2)
        self.add_input(layout, "Barcode:", 3)
        self.add_input(layout, "QIgG:", 4)
        self.add_input(layout, "QAlb:", 5)

        self.submit_button = QPushButton("Submit")
        self.submit_button.clicked.connect(self.save_data)
        layout.addWidget(self.submit_button, 6, 1)

        self.reset_button = QPushButton("Reset")
        self.reset_button.clicked.connect(self.reset_fields)
        layout.addWidget(self.reset_button, 6, 2)

        self.setLayout(layout)

    def add_input(self, layout, label_text, row):
        label = QLabel(label_text)
        entry = QLineEdit()
        layout.addWidget(label, row, 0, Qt.AlignmentFlag.AlignLeft)
        layout.addWidget(entry, row, 1)
        self.input_widgets[
            label_text
        ] = entry  # Store the QLineEdit widget in the dictionary

    def reset_fields(self):
        for widget in self.input_widgets.values():
            widget.clear()

    def save_data(self):
        # Get the entered values
        name = self.input_widgets["Name/Surname:"].text()
        age = self.input_widgets["Age:"].text()
        gender = self.input_widgets["Gender:"].text()
        barcode = self.input_widgets["Barcode:"].text()
        qigg = self.input_widgets["QIgG:"].text()
        qalb = self.input_widgets["QAlb:"].text()

        # Validate inputs
        if not all([name, age, gender, barcode, qigg, qalb]):
            QMessageBox.warning(
                self,
                "Validation Error",
                "Please fill in all fields.",
                QMessageBox.Ok,
            )
            return

        try:
            age = int(age)
            qigg = float(qigg) / 1000
            qalb = float(qalb) / 1000
        except ValueError:
            QMessageBox.warning(
                self,
                "Validation Error",
                "Invalid input for Age, QIgG, or QAlb.",
                QMessageBox.Ok,
            )
            return

        # Call the function to generate the Word document with information and the Reibergram plot
        folder_path = create_date_folder()
        self.generate_word_doc_with_info(
            qigg, qalb, name, age, gender, barcode, folder_path
        )

        # Reset the input fields
        self.reset_fields()
        QMessageBox.information(
            self,
            "Data Saved",
            "Data saved successfully.",
            QMessageBox.Ok,
        )

    def generate_word_doc_with_info(
        self, Qigg, Qalbumin, name, age, gender, barcode, folder_path
    ):
        doc_name = f"{barcode}.docx"
        plot_file = f"{barcode}.png"
        doc_path = os.path.join(folder_path, doc_name)

        # Create a Word document
        doc = Document()

        plot_reibergram(Qigg, Qalbumin, barcode)

        # Add collected information to the Word document
        info_text = (
            f"Name/Surname: {name}\nAge: {age}\nGender: {gender}\nBarcode: {barcode}"
        )

        doc.add_heading("Data Entries:", level=1)
        doc.add_paragraph(info_text)

        # Add the plot image to the Word document
        doc.add_picture(
            plot_file, width=Inches(5), height=Inches(5)
        )  # Adjust width and height as needed

        # Save the Word document
        doc.save(doc_path)

        # Delete redundant plot.png
        if os.path.exists(plot_file):
            os.remove(plot_file)

        # Close the plot to free up memory
        plt.close()


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()

        self.init_ui()

    def init_ui(self):
        self.setWindowTitle("REIBERGRAM by Sude Nur Cüre")
        self.setGeometry(100, 100, 800, 600)

        self.central_widget = QWidget(self)
        self.setCentralWidget(self.central_widget)

        layout = QVBoxLayout(self.central_widget)

        label = QLabel("Welcome to REIBERGRAM")
        label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(label)

        # Create a timer to close the main window and show the data entry window
        self.timer = QTimer(self)
        self.timer.timeout.connect(self.close_and_show_data_entry)
        self.timer.start(4000)  # Close after 4 seconds (4000 milliseconds)

    def close_and_show_data_entry(self):
        self.timer.stop()  # Stop the timer
        self.close()  # Close the main window
        self.data_entry_window = DataEntryWindow()
        self.data_entry_window.show()


if __name__ == "__main__":
    # Create a directory to store all documents
    if not os.path.exists("All Documents"):
        os.makedirs("All Documents")

    labels = ["Name/Surname", "Age", "Gender", "Barcode", "QIgG", "QAlb"]
    data_entries = []
    app = QApplication(sys.argv)
    main_window = MainWindow()
    main_window.show()
    sys.exit(app.exec_())
